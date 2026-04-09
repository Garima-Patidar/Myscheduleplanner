from __future__ import annotations

import csv
import os
import re
import sqlite3
from collections import defaultdict
from datetime import date, datetime, timedelta
from functools import wraps
from io import StringIO
import markdown
from flask import Flask, Response, flash, redirect, render_template, request, session, url_for
from werkzeug.utils import secure_filename
from collections import defaultdict

try:
    from openai import OpenAI
except Exception:
    OpenAI = None

try:
    from docx import Document
except Exception:
    Document = None

try:
    from pypdf import PdfReader
except Exception:
    PdfReader = None

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_PATH = os.path.join(BASE_DIR, 'planner.db')
UPLOAD_FOLDER = os.path.join(BASE_DIR, 'uploads')
ALLOWED_EXTENSIONS = {'txt', 'docx', 'pdf'}

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.secret_key = os.environ.get('FLASK_SECRET_KEY', 'my_schedule_planner_dev_key')

api_key = os.environ.get('OPENAI_API_KEY')
client = OpenAI(
    api_key="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",
    base_url="https://openrouter.ai/api/v1",
  
)


# ---------------------------
# Database helpers
# ---------------------------
def get_db() -> sqlite3.Connection:
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL;")
    conn.execute("PRAGMA busy_timeout = 30000;")
    return conn

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
def create_tables() -> None:
    conn = get_db()
    cur = conn.cursor()


    cur.execute(
        '''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            email TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL
        )
        '''
    )

    cur.execute(
        '''
        CREATE TABLE IF NOT EXISTS subjects (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            name TEXT NOT NULL,
            deadline TEXT NOT NULL,
            difficulty TEXT NOT NULL,
            FOREIGN KEY (user_id) REFERENCES users(id)
        )
        '''
    )

    cur.execute(
        '''
        CREATE TABLE IF NOT EXISTS topics (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            subject_id INTEGER NOT NULL,
            topic_name TEXT NOT NULL,
            status TEXT DEFAULT 'pending',
            FOREIGN KEY (subject_id) REFERENCES subjects(id)
        )
        '''
    )

    cur.execute(
        '''
        CREATE TABLE IF NOT EXISTS timetable (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            topic_id INTEGER UNIQUE NOT NULL,
            date TEXT NOT NULL,
            FOREIGN KEY (topic_id) REFERENCES topics(id)
        )
        '''
    )

    conn.commit()
    conn.close()

def ensure_topic_unit_column():
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute("ALTER TABLE topics ADD COLUMN unit_name TEXT DEFAULT 'General'")
        conn.commit()
    except:
        pass
    conn.close()

def add_unit_column_to_timetable():
    conn = get_db()
    cur = conn.cursor()
    try:
        cur.execute("ALTER TABLE timetable ADD COLUMN unit_name TEXT")
        conn.commit()
    except:
        pass
    conn.close()  
def normalize_topics(raw_topics: str) -> list[str]:
    items = [topic.strip() for topic in raw_topics.replace('\n', ',').split(',')]
    return [topic for topic in items if topic]

def allowed_file(filename: str) -> bool:
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def parse_topic_lines(raw_text: str) -> list[str]:
    cleaned = []
    seen = set()

    for line in raw_text.splitlines():
        value = re.sub(r'^[\s\-•*\d.)]+', '', line).strip(' :\t')
        if len(value) < 3:
            continue

        lower_val = value.lower()

        # common useless lines skip
        skip_words = [
            'unit', 'course outcomes', 'reference', 'text book',
            'books', 'practical', 'lab', 'marks', 'objective'
        ]
        if any(word in lower_val for word in skip_words) and len(value.split()) <= 6:
            continue

        if lower_val not in seen:
            cleaned.append(value)
            seen.add(lower_val)

    return cleaned



def extract_text_from_file(file_storage) -> str:
    filename = secure_filename(file_storage.filename)
    if not filename:
        return ''

    ext = filename.rsplit('.', 1)[1].lower()
    save_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    file_storage.save(save_path)

    if ext == 'txt':
        with open(save_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    if ext == 'docx' and Document:
        try:
            doc = Document(save_path)
            parts = []

            # normal paragraphs
            for p in doc.paragraphs:
                text = p.text.strip()
                if text:
                    parts.append(text)

            # tables bhi read karo
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            parts.append(cell_text)

            return '\n'.join(parts)
        except Exception as e:
            print("DOCX read error:", e)
            return ''

    if ext == 'pdf' and PdfReader:
        try:
            reader = PdfReader(save_path)
            parts = []
            for page in reader.pages:
                text = page.extract_text() or ''
                if text.strip():
                    parts.append(text)
            return '\n'.join(parts)
        except Exception as e:
            print("PDF read error:", e)
            return ''

    return ''

def split_large_topic(topic: str) -> list[str]:
    # long lines ko smaller chunks me todne ke liye
    separators = [';', ':', ' and ', ' & ', ',']
    results = [topic]

    for sep in separators:
        temp = []
        for item in results:
            if len(item) > 55 and sep in item:
                temp.extend([x.strip() for x in item.split(sep) if x.strip()])
            else:
                temp.append(item)
        results = temp

    final = []
    seen = set()
    for item in results:
        item = item.strip(' -•')
        if len(item) >= 3 and item.lower() not in seen:
            final.append(item)
            seen.add(item.lower())
    return final


def extract_topics_with_ai_or_rules(syllabus_text: str, subject_name: str) -> list[str]:
    if not syllabus_text.strip():
        return []

    # First try AI if available
    if client:
        try:
            response = client.chat.completions.create(
                model='gpt-4o-mini',
                messages=[
                    {
                        'role': 'system',
                        'content': (
                            'You are an academic syllabus parser. '
                            'Extract only clean study topics from a syllabus. '
                            'Return one topic per line. '
                            'Remove unit headings, numbering, duplicates, references, marks info, and filler text. '
                            'Break large syllabus statements into smaller meaningful study topics.'
                        )
                    },
                    {
                        'role': 'user',
                        'content': f"Subject: {subject_name}\n\nSyllabus:\n{syllabus_text}"
                    }
                ],
            )
            raw = response.choices[0].message.content
            ai_topics = parse_topic_lines(raw)

            expanded = []
            for t in ai_topics:
                expanded.extend(split_large_topic(t))

            return expanded[:80]
        except Exception:
            pass

    # fallback smart parsing
    base_topics = parse_topic_lines(syllabus_text)

    expanded = []
    for t in base_topics:
        expanded.extend(split_large_topic(t))

    # अगर file me proper lines hi nahi thi, then comma/newline based fallback
    if not expanded:
        expanded = normalize_topics(syllabus_text)

    return expanded[:80]
def days_left(deadline_text: str) -> int:
    deadline = datetime.strptime(deadline_text, '%Y-%m-%d').date()
    return (deadline - date.today()).days


def difficulty_weight(value: str) -> int:
    mapping = {'easy': 1, 'medium': 2, 'hard': 3}
    return mapping.get((value or '').strip().lower(), 2)


def login_required(view_func):
    @wraps(view_func)
    def wrapper(*args, **kwargs):
        if 'user_id' not in session:
            return redirect(url_for('login'))
        return view_func(*args, **kwargs)

    return wrapper


def get_subject_for_user(subject_id: int, user_id: int):
    conn = get_db()
    cur = conn.cursor()
    cur.execute('SELECT * FROM subjects WHERE id=? AND user_id=?', (subject_id, user_id))
    subject = cur.fetchone()
    conn.close()
    return subject


def fallback_notes(topic: str) -> str:
    return f"""Topic: {topic}

1. Definition
- {topic} ko simple language me samjho.
- Iske core concept aur purpose par focus karo.

2. Key Points
- Important terms identify karo.
- Exam me poochhe ja sakne wale concepts note karo.
- Short examples ke saath revise karo.

3. How to Study
- Pehle theory read karo.
- Phir 2-3 examples solve karo.
- End me ek short self-revision summary banao.

4. Quick Revision
- 3 keywords yaad rakho.
- 1 short definition likh kar practice karo.
- 1 practical example socho.

5. Exam Tip
- Is topic ka answer heading + points format me likhna best rahega.
"""



def extract_topics_unitwise(syllabus_text: str) -> list[dict]:
    lines = [line.strip() for line in syllabus_text.splitlines() if line.strip()]
    structured_topics = []
    current_unit = "General"

    for line in lines:
        clean = re.sub(r'\s+', ' ', line).strip()

        # unit heading detect
        if re.match(r'^(unit\s*\d+|module\s*\d+|chapter\s*\d+)', clean.lower()):
            current_unit = clean
            continue

        # junk skip
        if len(clean) < 3:
            continue

        junk_words = [
            "course outcomes", "reference", "text book", "books",
            "practical", "lab", "marks", "objective", "syllabus"
        ]
        if any(word in clean.lower() for word in junk_words) and len(clean.split()) <= 6:
            continue

        # split topics smartly
        parts = re.split(r',|;|\.| and ', clean)
        split_parts = [p.strip(" -•:\t") for p in parts if len(p.strip(" -•:\t")) > 2]

        for topic in split_parts:
            structured_topics.append({
                "unit": current_unit,
                "topic": topic
            })

    # remove duplicates
    final = []
    seen = set()
    for item in structured_topics:
        key = (item["unit"].lower(), item["topic"].lower())
        if key not in seen:
            final.append(item)
            seen.add(key)

    return final

def generate_ai_notes(topic: str, subject_name: str = "") -> str:
    if not client:
        return fallback_notes(topic)

    prompt = f"""
Create high-quality study notes for an engineering student.

Subject: {subject_name}
Topic: {topic}

Rules:
- Return well-structured markdown
- Use proper headings and subheadings
- Keep explanation simple and exam-friendly
- Use bullet points where helpful
- Add short examples
- Avoid one giant paragraph

Use this structure exactly:

# {topic}

## 1. Definition
## 2. Core Concept
## 3. Key Points
## 4. Real-life / Technical Example
## 5. Exam Answer (5 Marks)
## 6. Quick Revision Points
## 7. Self-Check Questions
"""

    try:
        response = client.chat.completions.create(
            model="openrouter/free",
            messages=[
                {"role": "system", "content": "You create detailed, structured study notes for engineering students."},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content
    except Exception:
        return fallback_notes(topic)

# ---------------------------
# Auth
# ---------------------------
@app.route('/', methods=['GET', 'POST'])
def login():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))

    error = None
    if request.method == 'POST':
        email = request.form.get('user_email', '').strip().lower()
        password = request.form.get('user_pass', '').strip()

        conn = get_db()
        cur = conn.cursor()
        cur.execute('SELECT * FROM users WHERE email=? AND password=?', (email, password))
        user = cur.fetchone()
        conn.close()

        if user:
            session['user_id'] = user['id']
            session['user_name'] = user['name']
            flash('Welcome back!', 'success')
            return redirect(url_for('dashboard'))

        error = 'Invalid email or password.'

    return render_template('login.html', error=error)


@app.route('/signup', methods=['GET', 'POST'])
def signup():
    if 'user_id' in session:
        return redirect(url_for('dashboard'))

    error = None
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        email = request.form.get('user_email', '').strip().lower()
        password = request.form.get('user_pass', '').strip()

        if not name or not email or not password:
            error = 'Please fill all fields.'
            return render_template('signup.html', error=error)

        conn = get_db()
        cur = conn.cursor()
        cur.execute('SELECT id FROM users WHERE email=?', (email,))
        existing_user = cur.fetchone()

        if existing_user:
            conn.close()
            error = 'Email already registered. Please log in.'
            return render_template('signup.html', error=error)

        cur.execute('INSERT INTO users (name, email, password) VALUES (?, ?, ?)', (name, email, password))
        conn.commit()
        conn.close()

        flash('Account created successfully. Please sign in.', 'success')
        return redirect(url_for('login'))

    return render_template('signup.html', error=error)


@app.route('/logout')
def logout():
    session.clear()
    flash('You have been logged out.', 'success')
    return redirect(url_for('login'))


# ---------------------------
# Dashboard
# ---------------------------
@app.route('/dashboard')
@login_required
def dashboard():
    user_id = session['user_id']
    today = date.today().strftime('%Y-%m-%d')

    conn = get_db()
    cur = conn.cursor()

    cur.execute('SELECT COUNT(*) FROM subjects WHERE user_id=?', (user_id,))
    total_subjects = cur.fetchone()[0]

    cur.execute(
        '''
        SELECT COUNT(*)
        FROM topics t
        JOIN subjects s ON t.subject_id = s.id
        WHERE s.user_id=?
        ''',
        (user_id,),
    )
    total_tasks = cur.fetchone()[0]

    cur.execute(
        '''
        SELECT COUNT(*)
        FROM topics t
        JOIN subjects s ON t.subject_id = s.id
        WHERE s.user_id=? AND t.status='done'
        ''',
        (user_id,),
    )
    completed = cur.fetchone()[0]

    cur.execute(
        '''
        SELECT t.id, t.topic_name, t.status, s.name AS subject_name, s.difficulty, s.deadline, tt.date
        FROM timetable tt
        JOIN topics t ON tt.topic_id = t.id
        JOIN subjects s ON t.subject_id = s.id
        WHERE s.user_id=? AND tt.date=? AND t.status!='done'
        ORDER BY lower(s.difficulty) DESC, t.topic_name ASC
        ''',
        (user_id, today),
    )
    today_tasks = cur.fetchall()

    cur.execute(
        '''
        SELECT s.name, s.deadline, COUNT(t.id) AS pending_topics
        FROM subjects s
        LEFT JOIN topics t ON t.subject_id = s.id AND t.status!='done'
        WHERE s.user_id=?
        GROUP BY s.id
        ORDER BY s.deadline ASC
        LIMIT 3
        ''',
        (user_id,),
    )
    upcoming_deadlines = cur.fetchall()
    conn.close()

    percent = int((completed / total_tasks) * 100) if total_tasks else 0
    pending = max(total_tasks - completed, 0)
    
    # ==== ADD THIS INSIDE dashboard() (IMPORTANT CHANGE) ====

    today_pending = len(today_tasks)

    if today_pending == 0 and total_tasks > 0:
        reminder_title = "🎉 All Tasks Completed"
        reminder_msg = "Amazing! Aaj ka pura study plan complete ho gaya 💪"
        reminder_type = "success"

    elif today_pending > 0:
        reminder_title = "⚠️ Pending Tasks"
        reminder_msg = f"You have {today_pending} tasks pending. Finish them before deadline."
        reminder_type = "warning"

    else:
        reminder_title = "🚀 Start Planning"
        reminder_msg = "Generate timetable and start your study journey."
        reminder_type = "info"
    
    return render_template(
        'dashboard.html',
        user_name=session.get('user_name', 'Student'),
        total_subjects=total_subjects,
        total_tasks=total_tasks,
        completed=completed,
        pending=pending,
        percent=percent,
        today=today,
        today_tasks=today_tasks,
        today_pending=len(today_tasks),
        upcoming_deadlines=upcoming_deadlines,
        reminder_title=reminder_title,
        reminder_msg=reminder_msg,
        reminder_type=reminder_type
    )


# ---------------------------
# Subjects
# ---------------------------
@app.route('/subjects')
@login_required
def subjects():
    conn = get_db()
    cur = conn.cursor()

    cur.execute(
        'SELECT * FROM subjects WHERE user_id=? ORDER BY deadline ASC, name ASC',
        (session['user_id'],)
    )
    subject_rows = cur.fetchall()

    subjects = []

    for row in subject_rows:
        cur.execute('SELECT * FROM topics WHERE subject_id=?', (row['id'],))
        topics = cur.fetchall()

        done = sum(1 for t in topics if t['status'] == 'done')
        total = len(topics)
        percent = int((done / total) * 100) if total > 0 else 0

        try:
            deadline_obj = datetime.strptime(row['deadline'], '%Y-%m-%d').date()
            days_left = max((deadline_obj - date.today()).days, 0)
        except Exception:
            days_left = 0

        subjects.append({
            'id': row['id'],
            'name': row['name'],
            'difficulty': row['difficulty'],
            'deadline': row['deadline'],
            'done': done,
            'total': total,
            'percent': percent,
            'days_left': days_left
        })

    conn.close()

    return render_template(
        'subjects.html',
        subjects=subjects,
        total_subjects=len(subjects)
    )

@app.route('/add_subject', methods=['GET', 'POST'])
@login_required
def add_subject():
    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        deadline = request.form.get('deadline', '').strip()
        difficulty = request.form.get('difficulty', 'Medium').strip().title()
        manual_topics = request.form.get('topics', '').strip()
        uploaded_file = request.files.get('syllabus_file')

        if not name or not deadline:
            flash('Subject name and deadline are required.', 'danger')
            return render_template('add_subject.html')

        structured_topics = []

        try:
            conn = get_db()
            cur = conn.cursor()

            if uploaded_file and uploaded_file.filename:
                if not allowed_file(uploaded_file.filename):
                    conn.close()
                    flash('Please upload only TXT, DOCX, or PDF file.', 'danger')
                    return render_template('add_subject.html')

                syllabus_text = extract_text_from_file(uploaded_file)

                if not syllabus_text.strip():
                    conn.close()
                    flash('Uploaded file could not be read properly. Please try another file.', 'danger')
                    return render_template('add_subject.html')

                structured_topics = extract_topics_unitwise(syllabus_text)

            elif manual_topics:
                manual_list = normalize_topics(manual_topics)
                structured_topics = [{"unit": "General", "topic": t} for t in manual_list]

            if not structured_topics:
                conn.close()
                flash('No topics could be extracted. Please upload a clearer syllabus file or add topics manually.', 'danger')
                return render_template('add_subject.html')

            cur.execute(
                'INSERT INTO subjects (user_id, name, deadline, difficulty) VALUES (?, ?, ?, ?)',
                (session['user_id'], name, deadline, difficulty),
            )
            subject_id = cur.lastrowid

            for item in structured_topics:
                cur.execute(
                    'INSERT INTO topics (subject_id, unit_name, topic_name) VALUES (?, ?, ?)',
                    (subject_id, item["unit"], item["topic"])
                )

            conn.commit()
            conn.close()

            flash(f'Subject added successfully with {len(structured_topics)} extracted topic(s).', 'success')
            return redirect(url_for('subjects'))

        except Exception as e:
            try:
                conn.close()
            except:
                pass
            flash(f'Error while saving subject: {str(e)}', 'danger')
            return render_template('add_subject.html')

    return render_template('add_subject.html')

@app.route('/edit_subject/<int:subject_id>', methods=['GET', 'POST'])
@login_required
def edit_subject(subject_id: int):
    user_id = session['user_id']
    conn = get_db()
    cur = conn.cursor()

    cur.execute('SELECT * FROM subjects WHERE id=? AND user_id=?', (subject_id, user_id))
    subject = cur.fetchone()
    if not subject:
        conn.close()
        flash('Subject not found.', 'danger')
        return redirect(url_for('subjects'))

    if request.method == 'POST':
        name = request.form.get('name', '').strip()
        deadline = request.form.get('deadline', '').strip()
        difficulty = request.form.get('difficulty', 'Medium').strip().title()

        cur.execute(
            'UPDATE subjects SET name=?, deadline=?, difficulty=? WHERE id=? AND user_id=?',
            (name, deadline, difficulty, subject_id, user_id),
        )
        conn.commit()
        conn.close()
        flash('Subject updated successfully.', 'success')
        return redirect(url_for('subjects'))

    conn.close()
    return render_template('edit_subject.html', subject=subject)


@app.route('/delete_subject/<int:subject_id>', methods=['POST'])
@login_required
def delete_subject(subject_id: int):
    user_id = session['user_id']
    conn = get_db()
    cur = conn.cursor()

    cur.execute('SELECT id FROM subjects WHERE id=? AND user_id=?', (subject_id, user_id))
    subject = cur.fetchone()
    if not subject:
        conn.close()
        flash('Subject not found.', 'danger')
        return redirect(url_for('subjects'))

    cur.execute('SELECT id FROM topics WHERE subject_id=?', (subject_id,))
    topic_ids = [row['id'] for row in cur.fetchall()]
    if topic_ids:
        placeholders = ','.join(['?'] * len(topic_ids))
        cur.execute(f'DELETE FROM timetable WHERE topic_id IN ({placeholders})', topic_ids)

    cur.execute('DELETE FROM topics WHERE subject_id=?', (subject_id,))
    cur.execute('DELETE FROM subjects WHERE id=? AND user_id=?', (subject_id, user_id))
    conn.commit()
    conn.close()

    flash('Subject deleted.', 'success')
    return redirect(url_for('subjects'))


@app.route('/subject/<int:subject_id>')
@login_required
def subject(subject_id: int):
    conn = get_db()
    cur = conn.cursor()

    cur.execute(
        'SELECT * FROM subjects WHERE id=? AND user_id=?',
        (subject_id, session['user_id'])
    )
    subject_row = cur.fetchone()

    if not subject_row:
        conn.close()
        flash('Subject not found.', 'danger')
        return redirect(url_for('subjects'))

    cur.execute(
        '''
        SELECT t.*, tt.date AS scheduled_date
        FROM topics t
        LEFT JOIN timetable tt ON tt.topic_id = t.id
        WHERE t.subject_id=?
        ORDER BY
            CASE
                WHEN t.unit_name IS NULL OR t.unit_name='' THEN 'ZZZ'
                ELSE t.unit_name
            END,
            t.id ASC
        ''',
        (subject_id,)
    )
    topics = cur.fetchall()
    conn.close()

    done_count = sum(1 for t in topics if t['status'] == 'done')
    total_topics = len(topics)
    pending_count = total_topics - done_count
    percent = int((done_count / total_topics) * 100) if total_topics else 0

    grouped_topics = defaultdict(list)
    for t in topics:
        unit = t['unit_name'] if t['unit_name'] else 'General'
        grouped_topics[unit].append(t)

    return render_template(
        'subject.html',
        subject=subject_row,
        grouped_topics=dict(grouped_topics),
        percent=percent,
        total_topics=total_topics,
        done_count=done_count,
        pending_count=pending_count
    )


# ---------------------------
# Notes
# ---------------------------
@app.route('/generate_notes/<int:topic_id>')
@login_required
def generate_notes(topic_id: int):
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        '''
        SELECT t.topic_name, s.name AS subject_name
        FROM topics t
        JOIN subjects s ON t.subject_id = s.id
        WHERE t.id=? AND s.user_id=?
        ''',
        (topic_id, session['user_id']),
    )
    topic = cur.fetchone()
    conn.close()

    if not topic:
        flash('Topic not found.', 'danger')
        return redirect(url_for('subjects'))

    notes_raw = generate_ai_notes(topic['topic_name'], topic['subject_name'])
    notes_html = markdown.markdown(
        notes_raw,
        extensions=['extra', 'nl2br', 'sane_lists']
    )

    return render_template(
        'notes.html',
        topic=topic['topic_name'],
        subject_name=topic['subject_name'],
        notes=notes_html
    )
 
@app.route('/study_chat', methods=['GET', 'POST'])
@login_required
def study_chat():
    conn = get_db()
    cur = conn.cursor()
    cur.execute('SELECT id, name FROM subjects WHERE user_id=? ORDER BY name ASC', (session['user_id'],))
    subjects = cur.fetchall()
    conn.close()

    if 'chat_history' not in session:
        session['chat_history'] = []

    selected_subject = request.form.get('subject_name', '') if request.method == 'POST' else ''
    user_question = ''

    if request.method == 'POST':
        user_question = request.form.get('question', '').strip()
        subject_name = selected_subject or "General"

        if not user_question:
            flash('Please enter a question.', 'danger')
            return render_template(
                'study_chat.html',
                subjects=subjects,
                selected_subject=selected_subject,
                chat_history=session['chat_history']
            )

        session['chat_history'].append({
            'role': 'user',
            'content': user_question
        })

        if not client:
            answer_html = """
            <h3>AI unavailable</h3>
            <p>API client configured nahi hai. Please check your API key setup.</p>
            """
        else:
            try:
                prompt = f"""
You are a smart study assistant for engineering students.

Subject: {subject_name}
Question: {user_question}

Rules:
- Explain in clean readable format
- Use short headings
- Use bullet points where helpful
- Keep language simple
- Make it feel like study notes + explanation
- Avoid very long paragraphs
- Give practical examples where possible
"""

                response = client.chat.completions.create(
                    model="openrouter/free",
                    messages=[
                        {"role": "system", "content": "You are a helpful study assistant."},
                        {"role": "user", "content": prompt}
                    ]
                )

                answer_raw = response.choices[0].message.content
                answer_html = markdown.markdown(
                    answer_raw,
                    extensions=['extra', 'nl2br', 'sane_lists']
                )

            except Exception as e:
                answer_html = f"""
                <h3>AI response issue</h3>
                <p>{str(e)}</p>
                """

        session['chat_history'].append({
            'role': 'assistant',
            'content': answer_html,
            'subject': subject_name
        })
        session.modified = True

        return render_template(
            'study_chat.html',
            subjects=subjects,
            selected_subject=selected_subject,
            chat_history=session['chat_history']
        )

    return render_template(
        'study_chat.html',
        subjects=subjects,
        selected_subject=selected_subject,
        chat_history=session['chat_history']
    )
@app.route('/clear_study_chat', methods=['POST'])
@login_required
def clear_study_chat():
    session['chat_history'] = []
    session.modified = True
    return redirect(url_for('study_chat'))
# ---------------------------
# Timetable
# ---------------------------
def build_timetable_for_user(user_id: int) -> None:
    conn = get_db()
    cur = conn.cursor()

    cur.execute(
        'DELETE FROM timetable WHERE topic_id IN (SELECT t.id FROM topics t JOIN subjects s ON t.subject_id=s.id WHERE s.user_id=?)',
        (user_id,)
    )

    cur.execute('''
        SELECT
            t.id,
            t.topic_name,
            t.unit_name,
            s.name AS subject_name,
            s.deadline,
            s.difficulty
        FROM topics t
        JOIN subjects s ON t.subject_id = s.id
        WHERE s.user_id=? AND t.status!='done'
        ORDER BY s.deadline ASC, t.unit_name ASC, t.id ASC
    ''', (user_id,))
    rows = cur.fetchall()

    today = date.today()

    # Group by subject + unit
    grouped = defaultdict(list)
    for row in rows:
        deadline = datetime.strptime(row['deadline'], '%Y-%m-%d').date()
        grouped[(row['subject_name'], row['unit_name'], deadline, row['difficulty'])].append(row)

    daily_load = defaultdict(int)

    for (subject_name, unit_name, deadline, difficulty), topics in grouped.items():
        days_available = max((deadline - today).days + 1, 1)

        # छोटी topics ko 3-4 ek din me club karna
        total_topics = len(topics)

        if total_topics <= days_available:
            chunk_size = 1
        elif total_topics <= days_available * 2:
            chunk_size = 2
        elif total_topics <= days_available * 3:
            chunk_size = 3
        else:
            chunk_size = 4

        chunks = [topics[i:i + chunk_size] for i in range(0, len(topics), chunk_size)]

        current_day = today

        for chunk in chunks:
            while daily_load[current_day] >= 4 and current_day <= deadline:
                current_day += timedelta(days=1)

            if current_day > deadline:
                # deadline ke baad nearest available
                current_day = deadline
                while daily_load[current_day] >= 4:
                    current_day += timedelta(days=1)

            for topic in chunk:
                cur.execute(
                    'INSERT OR REPLACE INTO timetable (topic_id, date) VALUES (?, ?)',
                    (topic['id'], current_day.strftime('%Y-%m-%d'))
                )

            daily_load[current_day] += 1
            current_day += timedelta(days=1)

    conn.commit()
    conn.close()
   
    

@app.route('/generate_timetable')
@login_required
def generate_timetable():
    conn = get_db()
    cur = conn.cursor()

    # get subjects
    cur.execute("SELECT * FROM subjects WHERE user_id=?", (session['user_id'],))
    subjects = cur.fetchall()

    today = date.today()

    # clear old timetable
    cur.execute("DELETE FROM timetable")

    for subject in subjects:

        cur.execute("""
            SELECT * FROM topics 
            WHERE subject_id=? AND status!='done'
            ORDER BY unit_name, id
        """, (subject['id'],))
        topics = cur.fetchall()

        if not topics:
            continue

        deadline = datetime.strptime(subject['deadline'], "%Y-%m-%d").date()
        total_days = max((deadline - today).days, 1)

        # 🔥 STEP 1: group topics
        grouped = []
        temp_group = []

        for t in topics:
            words = len(t['topic_name'].split())

            if words <= 3:
                temp_group.append(t)

                if len(temp_group) == 3:
                    grouped.append(temp_group)
                    temp_group = []
            else:
                if temp_group:
                    grouped.append(temp_group)
                    temp_group = []
                grouped.append([t])

        if temp_group:
            grouped.append(temp_group)

        # 🔥 STEP 2: distribute across days
        per_day = max(1, len(grouped) // total_days)

        day_counter = 0
        current_date = today

        for group in grouped:
            for topic in group:
                cur.execute("""
                    INSERT INTO timetable (topic_id, date, unit_name)
                    VALUES (?, ?, ?)
                """, (
                    topic['id'],
                    current_date.strftime("%Y-%m-%d"),
                    topic['unit_name']
                ))

            day_counter += 1

            if day_counter >= per_day:
                current_date += timedelta(days=1)
                day_counter = 0

    conn.commit()
    conn.close()

    flash("Smart timetable generated successfully!", "success")
    return redirect(url_for('dashboard'))


@app.route('/timetable')
@login_required
def timetable():
    user_id = session['user_id']
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        '''
        SELECT tt.date, t.id AS topic_id, t.topic_name, t.status, s.name AS subject_name, s.difficulty, s.deadline
        FROM timetable tt
        JOIN topics t ON tt.topic_id = t.id
        JOIN subjects s ON t.subject_id = s.id
        WHERE s.user_id=?
        ORDER BY tt.date ASC, t.status ASC, s.name ASC
        ''',
        (user_id,),
    )
    rows = cur.fetchall()
    conn.close()

    grouped = defaultdict(list)
    for row in rows:
        grouped[row['date']].append(row)

    return render_template('timetable.html', grouped_timetable=dict(grouped), today=date.today().strftime('%Y-%m-%d'))

@app.route('/analytics')
@login_required
def analytics():
    user_id = session['user_id']
    conn = get_db()
    cur = conn.cursor()

    cur.execute('SELECT * FROM subjects WHERE user_id=? ORDER BY name ASC', (user_id,))
    subject_rows = cur.fetchall()

    data = []
    total_topics = 0
    total_done = 0

    for s in subject_rows:
        cur.execute('SELECT * FROM topics WHERE subject_id=?', (s['id'],))
        topics = cur.fetchall()

        done = sum(1 for t in topics if t['status'] == 'done')
        total = len(topics)
        pending = total - done
        percent = int((done / total) * 100) if total > 0 else 0

        if percent >= 75:
            risk_status = 'On Track'
            risk_class = 'track'
        elif percent >= 40:
            risk_status = 'Moderate'
            risk_class = 'moderate'
        else:
            risk_status = 'Critical'
            risk_class = 'critical'

        total_topics += total
        total_done += done

        data.append({
            'id': s['id'],
            'name': s['name'],
            'total': total,
            'done': done,
            'pending': pending,
            'percent': percent,
            'difficulty': s['difficulty'],
            'deadline': s['deadline'],
            'risk_status': risk_status,
            'risk_class': risk_class
        })

    conn.close()

    total_pending = total_topics - total_done
    overall = int((total_done / total_topics) * 100) if total_topics > 0 else 0

    best_subject = max(data, key=lambda x: x['percent']) if data else None
    risk_subject = max(data, key=lambda x: x['pending']) if data else None

    return render_template(
        'analytics.html',
        data=data,
        total_topics=total_topics,
        total_done=total_done,
        total_pending=total_pending,
        overall=overall,
        best_subject=best_subject,
        risk_subject=risk_subject
    )

@app.route('/download_report')
@login_required
def download_report():
    user_id = session['user_id']
    conn = get_db()
    cur = conn.cursor()

    cur.execute('SELECT * FROM subjects WHERE user_id=? ORDER BY name ASC', (user_id,))
    subjects = cur.fetchall()

    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(['Subject', 'Total Topics', 'Completed', 'Pending', 'Progress %', 'Difficulty', 'Deadline'])

    for s in subjects:
        cur.execute('SELECT * FROM topics WHERE subject_id=?', (s['id'],))
        topics = cur.fetchall()

        done = sum(1 for t in topics if t['status'] == 'done')
        total = len(topics)
        pending = total - done
        percent = int((done / total) * 100) if total > 0 else 0

        writer.writerow([
            s['name'],
            total,
            done,
            pending,
            percent,
            s['difficulty'],
            s['deadline']
        ])

    conn.close()
    output.seek(0)

    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': 'attachment; filename=subject_progress_report.csv'}
    )

@app.route('/view_report')
@login_required
def view_report():
    user_id = session['user_id']
    conn = get_db()
    cur = conn.cursor()

    cur.execute('SELECT * FROM subjects WHERE user_id=? ORDER BY name ASC', (user_id,))
    subjects = cur.fetchall()

    report_data = []
    total_topics = 0
    total_done = 0

    for s in subjects:
        cur.execute(
            '''
            SELECT t.topic_name, t.status, tt.date as scheduled_date
            FROM topics t
            LEFT JOIN timetable tt ON tt.topic_id = t.id
            WHERE t.subject_id=?
            ORDER BY t.topic_name ASC
            ''',
            (s['id'],)
        )
        topics = cur.fetchall()

        done = sum(1 for t in topics if t['status'] == 'done')
        total = len(topics)
        pending = total - done
        percent = int((done / total) * 100) if total > 0 else 0

        total_topics += total
        total_done += done

        report_data.append({
            'name': s['name'],
            'difficulty': s['difficulty'],
            'deadline': s['deadline'],
            'total': total,
            'done': done,
            'pending': pending,
            'percent': percent,
            'topics': topics
        })

    conn.close()

    total_pending = total_topics - total_done
    overall = int((total_done / total_topics) * 100) if total_topics > 0 else 0

    return render_template(
        'view_report.html',
        report_data=report_data,
        total_topics=total_topics,
        total_done=total_done,
        total_pending=total_pending,
        overall=overall,
        generated_on=datetime.now().strftime('%d-%m-%Y %I:%M %p'),
        student_name=session.get('user_name', 'Student')
    )


@app.route('/download_subject_log/<int:subject_id>')
@login_required
def download_subject_log(subject_id):
    user_id = session['user_id']
    conn = get_db()
    cur = conn.cursor()

    cur.execute('SELECT * FROM subjects WHERE id=? AND user_id=?', (subject_id, user_id))
    subject = cur.fetchone()

    if not subject:
        conn.close()
        flash('Subject not found.', 'danger')
        return redirect(url_for('analytics'))

    cur.execute(
        '''
        SELECT t.topic_name, t.status, tt.date as scheduled_date
        FROM topics t
        LEFT JOIN timetable tt ON tt.topic_id = t.id
        WHERE t.subject_id=?
        ORDER BY t.topic_name ASC
        ''',
        (subject_id,)
    )
    topics = cur.fetchall()
    conn.close()

    output = StPringIO()
    writer = csv.writer(output)
    writer.writerow(['Subject', subject['name']])
    writer.writerow(['Difficulty', subject['difficulty']])
    writer.writerow(['Deadline', subject['deadline']])
    writer.writerow([])
    writer.writerow(['Topic Name', 'Status', 'Scheduled Date'])

    for t in topics:
        writer.writerow([t['topic_name'], t['status'], t['scheduled_date'] or 'Not Scheduled'])

    output.seek(0)

    safe_name = subject['name'].replace(' ', '_').lower()
    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': f'attachment; filename={safe_name}_topic_log.csv'}
    )

@app.route('/mark_done/<int:topic_id>', methods=['POST'])
@login_required
def mark_done(topic_id: int):
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        '''
        UPDATE topics
        SET status='done'
        WHERE id=? AND subject_id IN (SELECT id FROM subjects WHERE user_id=?)
        ''',
        (topic_id, session['user_id']),
    )
    cur.execute('DELETE FROM timetable WHERE topic_id=?', (topic_id,))
    conn.commit()
    conn.close()

    flash('Great! Topic marked as completed.', 'success')
    return redirect(request.referrer or url_for('dashboard'))


@app.route('/skip_task/<int:topic_id>', methods=['POST'])
@login_required
def skip_task(topic_id: int):
    conn = get_db()
    cur = conn.cursor()
    cur.execute(
        '''
        SELECT tt.date, s.deadline
        FROM timetable tt
        JOIN topics t ON tt.topic_id = t.id
        JOIN subjects s ON t.subject_id = s.id
        WHERE tt.topic_id=? AND s.user_id=?
        ''',
        (topic_id, session['user_id']),
    )
    row = cur.fetchone()

    if row:
        old_date = datetime.strptime(row['date'], '%Y-%m-%d').date()
        subject_deadline = datetime.strptime(row['deadline'], '%Y-%m-%d').date()
        new_date = old_date + timedelta(days=1)
        if new_date > subject_deadline:
            new_date = subject_deadline

        cur.execute('UPDATE timetable SET date=? WHERE topic_id=?', (new_date.strftime('%Y-%m-%d'), topic_id))
        conn.commit()
        flash('Task rescheduled to the next available day.', 'success')
    else:
        flash('Task schedule not found.', 'danger')

    conn.close()
    return redirect(request.referrer or url_for('dashboard'))


@app.context_processor
def inject_globals():
    return {'project_name': 'My Schedule Planner'}


if __name__ == '__main__':
    create_tables()
    ensure_topic_unit_column()
    add_unit_column_to_timetable()
    app.run(debug=True, use_reloader=False)
